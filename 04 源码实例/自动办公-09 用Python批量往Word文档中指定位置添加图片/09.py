#!/usr/bin/env python
# coding: utf-8

# 公司每周都要搞6S稽核（在制造业工作过的同学，应该知道)。 稽核会发现一些问题，稽核人会记录下问题，并拍照片。等稽核完成，需要将稽核的记录和图片放入word文档，并发给相关部门回复改善对策。文字的部分还好，本来就要手动录入，但遇到图片，就难搞了。将图片复制粘贴进去后，是这样的：
# ![](images\problem.png)
# 可见，把整个的表格都撑变形了。然后，需要手工将其拖拽，缩小到如下这样的适合尺寸。一般每次稽核，都会发现40~60个问题，手动一个一个地调整，时间就被浪费在这机械的重复中了。人生苦短，还是上Python吧。
# ![](images\after.png)
# 我们可以把序号、问题描述和责任部门都先key进Excel表格中，然后再用Python统一来填写。

# In[44]:


#提取Excel表中的数据
from openpyxl import load_workbook #用于读取Excel中的信息
wb = load_workbook('数据.xlsx')
ws = wb.active
data=[]
for row in range(2,ws.max_row+1):
    number = ws['A' + str(row)].value
    problem = ws['B' + str(row)].value
    owner = ws['C' + str(row)].value
    info_list = [number,problem,owner]
    data.append(info_list)


# In[45]:


data[:5]


# In[46]:


len(data)


# 在写入数据到Word表格之前，需要获取图片所在路径，并按“修改时间”升序排序。因为默认是按数字及字母顺序排序（如下图），这样会打乱图片顺序，导致插入图片后，会与文字描述不符。通过`os.listdir()`获得图片文件夹内的所有图片的文件名，然后使用`sort`对所有图片进行排序，排序依据为图片的修改时间。`os.path.getmtime`可获取文件修改时间。
# ![](images\wrong_seq.png)

# `list_p = [path+"\\"+i for i in os.listdir(path)]` 是列表解析式的写法，可让程序显得简洁，并提高运行速度。它跟如下写法的结果是相同的。
# 
# ```
# list_p=[]
# for i in os.listdir(path):
#     list_p.append(path+"\\"+i)
# ```

# 所有图片的路径获取完成后，使用`sort`方法重新排序，排序依据为该文件的修改时间。通过`os.path.getmtime()`来获得文件的修改时间。此处使用了匿名函数`lambda`的方式`list_p.sort(key=lambda path: os.path.getmtime(path))`，以让程序显得简洁。它跟如下写法是相同的效果。
# ```
# def Get_time(path):
#     return os.path.getmtime(path) #获取文件修改时间
#     
# list_p.sort(key=Get_time) #以文件修改时间为依据升序排序
# ```

# 图片顺序处理好后，就要新建足够的单元格，准备填数据和图片了。依据数据的个数`len(data)`来新增的表格行数，因为模板中已经有一行空白表格，所以此处要减掉1`len(data)-1`。通过`add_row()`在下面插入一行。

# 然后用`for`循环，往单元格中填入数据。数据填完后，往第3列中插入图片。由于只有文字块`run`才可使用`add_picture`方法，所以要先新增文字块，存入`run`变量，通过在`add_picture`中传入图片路径，插入图片。然后强行设置图片的高度和宽度，以匹配模板中图片单元格的尺寸。此处，模板中单元格的高度是4.42厘米，宽度是6.25厘米。我们设置图片的高度为4.4厘米，宽度6.2厘米，几乎刚好放进单元格。
# ![](images\size.png)

# In[74]:


from docx import Document
from docx.shared import Cm
import  os
doc = Document("6S稽查问题模板.docx")

#将图片按修改时间排序(这样才能与图片的描述一致)，将路径存入列表，以便后面逐个插入图片时调用
path = "6s_pictures"
list_p = [path+"\\"+i for i in os.listdir(path)] #获取图片的文件名,并拼接完整路径
list_p.sort(key=lambda path: os.path.getmtime(path)) #将列表中的文件按其修改时间排序，os.path.getmtime() 函数是获取文件最后修改时间

table = doc.tables[0] #已确定是第一个表格，其索引是0

#增加需要的行，以便足够填入数据
for i in range(len(data)-1):
    table.add_row()

#写入数据及图片
for row in range(1,len(data)+1):
    table.cell(row,0).text = str(data[row-1][0]) #往第1列写入序号
    table.cell(row,1).text = data[row-1][1] #往第2列写入问题描述
    table.cell(row,3).text = data[row-1][2] #往第4列写入责任部门
    
    #插入图片并调整图片的高度和宽度，以适合模板中的单元格尺寸
    run = table.cell(row,2).paragraphs[0].add_run() #新增一个文字块
    picture = run.add_picture(list_p[row-1]) #插入图片
    picture.height = Cm(4.4)  #设置图片高度
    picture.width = Cm(6.2) #设置图片宽度

doc.save("6S稽查问题.docx")


# 结果如下。这个结果是不完美的，因为图片被强行设定了高度和宽度后，会变形，像哈哈镜里看到的一样。需要根据其本身的高度和宽度按比例设置新的高度和宽度才能解决这个问题（具体操作后续再分享吧）。但如果图片只是用于示意，这样的效果也足够了。
# ![](images\result.png)

# In[ ]:




