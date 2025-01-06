#!/usr/bin/env python
# coding: utf-8

# 公司的很多会议通知文件，除了“第几次”啊、时间啊不同，主要内容基本一致。对于批量的操作，用Python来快速替换这些字符，效率贼高。于是调出之前定义好的python字符替换函数，一顿操作猛如虎，结果目瞪狗呆了。纳尼？替换前和替换后就是双胞胎，长相分毫不差啊。

# In[4]:


#定义函数替换文字块中的字符
def info_update(doc,old_info, new_info): 
    import docx
    for para in doc.paragraphs: #遍历段落
        for run in para.runs: #遍历文字块
            run.text = run.text.replace(old_info, new_info) #替换

import docx
doc = docx.Document('替换前.docx')
info_update(doc,"第四次", "第五次")
info_update(doc,"2019", "2020")
info_update(doc,"18", "10")
doc.save('替换后.docx')


# 到底发生了什么事？函数是没问题的，那查一下目标word文件。果然发现了幺蛾子。原来需要替换的“第四次”、“2019”、“18”三个字符串居然都不在同一个文字块（即程序中的`run`）。其中“第四次”作为跨栏高手，横跨3,4,5三个文字块，“2019”腿短一点，跨第6,7文字块，最小的最没跨栏可能的“18”居然也跨了第8,9文字块。难怪一个也没替换完成，因为在同一个文字块根本都找不到它们哪。Python大喊冤枉。
# ![](images\runs.png)

# In[3]:


import docx
doc = docx.Document('替换前.docx')
for para in doc.paragraphs: 
    for run in para.runs:
        print(run.text)


# 对于word文档的处理，文字块是个很头疼的问题，稍微一点调整，就会打破文字块，导致文字“跨栏”。那怎么办呢？一大堆文档，总不能放弃Python，改用手工替换吧？办法还是有的，既然在文字块中替换不行，那在段落中替换行不行呢？试试看吧。先按段落打印一下文档。哟呵，完整的两个段落，看起来规规整整的嘛，有希望。

# In[5]:


#按段落查看文字
doc = docx.Document('替换前.docx')
for para in doc.paragraphs: 
    print(para.text)


# 于是乎，充满期待地开始按段落替换文字。

# In[9]:


#定义函数替换段落中的字符
def info_update_para(doc,old_info, new_info): 
    import docx
    for para in doc.paragraphs: #遍历段落
        para.text = para.text.replace(old_info, new_info) #替换

doc = docx.Document('替换前.docx')
info_update_para(doc,"第四次", "第五次")
info_update_para(doc,"2019", "2020")
info_update_para(doc,"18", "10")
doc.save('替换后_段落.docx')


# 还没来得及眨一下眼睛，程序运行完成了，迫不及待打开文件。
# ![](images\replace.png)
# 毫无疑问，替换那是相当的成功。然鹅，好像有点点不对。哪里不对呢？呕，我的格式呢？我的微软雅黑呢，我的加粗呢，我标题怎么廋了呢？原来只有在文字块中替换，才不会影响到原文档格式。用段落替换，所有的格式都被打回原形，宋体五号。泪崩中......请勿扰。总不能替换完后，又去一个一个打开原文档，一个一个修改回原来设定的格式吧？回答肯定是“不用”，可以在段落替换文字后重新设定格式。收拾收拾，改程序去。

# In[13]:


import docx
from docx.shared import Pt #用于设定字体大小（磅值）
from docx.oxml.ns import qn #用于应用中文字体

def Info_update(doc,old_info, new_info):
    for para in doc.paragraphs: 
        para.text = para.text.replace(old_info, new_info)
        
    #设置第一段（标题）的文字格式   
    for run in doc.paragraphs[0].runs:
        run.font.size = Pt(14) #文字大小磅值
        run.bold = True #加粗
        run.font.name = "微软雅黑" #字体选择
        #中文字体应用，固定写法
        r = run._element.rPr.rFonts #字体，固定写法
        r.set(qn("w:eastAsia"),"微软雅黑") #字体

    
    #设置第二及后续段落的文字格式
    for para in doc.paragraphs[1:]: 
        for run in para.runs:
            run.font.size = Pt(12)  #文字大小
            run.bold = False #不加粗
            run.font.name = "微软雅黑" #字体选择
            #中文字体应用，固定写法
            r = run._element.rPr.rFonts 
            r.set(qn("w:eastAsia"),"微软雅黑") 
            
doc = docx.Document('替换前.docx')
Info_update(doc,"第四次", "第五次")
Info_update(doc,"2019", "2020")
Info_update(doc,"18", "10")
doc.save('替换后_设置格式.docx')


# 以上，用于文字替换的程序保持不变。新增了文字格式设定的程序，包括文字大小（磅值）的设置，是否加粗，字体及中文字体应用。其中中文字体应用是固定写法，照抄即可。不过字体名字需要与`run.font.name`中的一致。在替换后按照原文件重新设置了段落的格式，终于实现了完美替换。当然，如果文档的段落较多，且格式种类也多的话，需要更多的程序进行设置。不过相对手工去word文档中设置，效率还是要高不止那么一点点的。
