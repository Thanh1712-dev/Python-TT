#!/usr/bin/env python
# coding: utf-8

# 公司的部分文件分为“内部”和“外部”。正常情况下，这个标识在文件名及文件首页左上角都有标注。然鹅，有时候操作一下，忘记改了，或者忘记标注了。就需要回头去一个一个地整理，非常麻烦。这种重复的，繁杂的操作，尽快冲Python来，人类还是应该多做做其它更有意思的事。
# <br/>
# <br/>共有如下7种情况，其对应的处理方式如下表格所示。

# ![](images\require.png)

# 在处理文档之前，发现还有很多是`.doc`为后缀的文件。由于`docx`库不支持这些文件，所以需要先将它们批量转换成`.docx`文件。程序如下。细节请参考实例15 https://zhuanlan.zhihu.com/p/64189783。 在完成转换后，用`os.remove()`将老破小的`doc`文件全部删除。

# In[8]:


#将文件夹内的所有doc转存为docx文件
import os #用于获取目标文件所在路径
path=os.getcwd()+"\\文件\\" # 文件夹绝对路径
files=[]
for file in os.listdir(path):
    if file.endswith(".doc"): #排除文件夹内的其它干扰文件，只获取".doc"后缀的word文件
        files.append(path+file) 
        
from win32com import client as wc #导入模块
word = wc.Dispatch("Word.Application") # 打开word应用程序
for file in files:
    doc = word.Documents.Open(file) #打开word文件
    doc.SaveAs("{}x".format(file), 12)#另存为后缀为".docx"的文件，其中参数12指docx文件
    doc.Close() #关闭原来word文件
word.Quit()

for file in files: #删除doc文件
    os.remove(file)


# 这下整个世界清静了，满满的都是`docx`文件。下面将文件夹内的所有`docx`文件路径全部获取，存入列表`docx_files`。

# In[ ]:


#获取所有docx文件路径
docx_files=[]
for file in os.listdir(path):
    if file.endswith(".docx"): #排除文件夹内的其它干扰文件，只获取".doc"后缀的word文件
        docx_files.append(path+file) 
docx_files


# 然后开始憋最后的大招，写整理文件的程序了。先导入相关的库。建一个计数器`counter`，令其初始值为0。这个计数器用于记录有多少文件没有被处理，然后用文件的总数减去它，就能得到处理过的文件数量。为什么不直接记录处理过的文件的数量呢？因为处理的操作太多，比较混乱，计数麻烦，所以曲线救国更方便。
# <br/>
# <br/>然后遍历待检查的所有文件，逐个进行处理。先用`split()`按"."拆分出文件头`file_head`，并用`rstrip()`去除尾部可能的空格；按"\\"拆分出文件尾`file_tail`。比如“C:\\Users\\文件\\测试.docx”拆出的头是“C:\\Users\\文件\\测试”，尾是“测试.docx”。这个后面重命名文件的时候会用到。
# <br/>
# <br/>将要用到的标记词“公开”和“内部”放入列表`mark_words`以便后续调用。因为文件头中有标记词的话，倒数第三个加第二个就是，因此按此取数存入字符串`mark_fileName`。然后使用`docx.Document()`打开该文件，提出首段`paragraphs[0]`中的文本，存入`mark_doc`。如果文档中有标记词的话，`mark_doc`的值就应该是“公开”或“内部”，否则就是其他字符了。
# <br/>
# <br/>标记词提取好了，就开始比对了，然后根据比对的结果采取对应的操作。以文档中是否有关键词，分为两种大的情况。
# <br/>
# <br/>如果文档中有标记词，再判断文件名中有无标记词，如果有的话，则将其跟文档中的进行比对，如果二者一致，则计数器加1，并通过`pass`跳过后面的所有判断，回到`for`循环，进行下一个文件的检查。如果二者标记词不同，则按照文档中的标记词重命名文件。如果文件名中无标记词，则给文件名加上标记词。
# <br/>
# <br/>如果文档中无标记词，则直接在文档首段前插入一段，并写入“公开”二次，并设置字体为黑体，大小为16磅，并保存文件。然后再判断文件名中是否有标记词，有的话再判断是否为“公开”，若是，则直接`pass`；若不是，则替换成“公开”。如果文件名中无标记词，则按“公开”标记词重命名文件。
# <br/>
# <br/>以上，只要有一步有对文件或文件名有操作，则通过`print`显示操作内容，以便我们知道做了哪些更改。最后显示检查过的文件和处理过的文件的个数。

# In[10]:


import docx
from docx.shared import Pt #用于设定字体大小（磅值）
from docx.oxml.ns import qn #用于应用中文字体

counter = 0 #计数器，用于记录有多少文件没被处理

for file in docx_files:    
    file_head = file.split('.')[0].rstrip() #文件名头，类似这样的“C:\\Users\\文件\\测试”
    file_tail = file.split('\\')[-1] #文件名尾，类似这样“测试.docx”
    mark_words = ["公开", "内部"]
    mark_fileName = file_head[-3:-1]#文件名中倒数第2,3个文字
    
    doc = docx.Document(file)
    mark_doc = doc.paragraphs[0].text #文件中首段文字
    #比对标记词
    if mark_doc in mark_words: #判断文件中有无标记
        if mark_fileName in mark_words:#判断文件名中有无标记
            if mark_doc == mark_fileName:#如果二者标记相同
                counter += 1
                pass
            else: #二者标记不同            
                os.rename(file, file.replace(mark_fileName, mark_doc)) #重命名文件
                print(f"【{file_tail}】文件名重命名标识为【{mark_doc}】")
        else:
            os.rename(file, f"{file_head}（{mark_doc}）.docx")#文件名中无标记，则加标记
            print(f"【{file_tail}】文件名增加标识为【{mark_doc}】")
    else:
        #文中无标记，则在首段前插入一段，写入标记
        p = doc.paragraphs[0] 
        pNew = p.insert_paragraph_before() 
        run = pNew.add_run(mark_words[0]) #写为“公开”
         #字体设置
        run.font.size = Pt(16)
        run.font.name = "黑体"
        r = run._element.rPr.rFonts
        r.set(qn("w:eastAsia"),"黑体") 
        
        doc.save(file)
        print(f"【{file_tail}】内容增加标识为【{mark_words[0]}】")
        
        if mark_fileName in mark_words:#查看文件名中是否有标记
            if mark_fileName == mark_words[0]: #标记是否为“公开”
                pass
            else: #标记不是“公开”则替换           
                os.rename(file, file.replace(mark_fileName,mark_words[0])) #重命名文件
                print(f"【{file_tail}】文件名重命名标识为【{mark_words[0]}】")
        else:
            os.rename(file, f"{file_head}（{mark_words[0]}）.docx")#文件名中无标记，则加标记
            print(f"【{file_tail}】文件名增加标识为【{mark_words[0]}】") 
            
print(f"完成！共检查{len(docx_files)}个文件,处理了 {len(docx_files)-counter} 个文件。")


# 到此，程序圆满完成任务，运行前和运行后的文件如下所示。可以看到，“测试.doc”文件被转换成了“.docx”文件，且按照文件里面的标记词“内部”重新命名了。其它几个测试文件也做了相应的修正。

# ![](images\result.png)
